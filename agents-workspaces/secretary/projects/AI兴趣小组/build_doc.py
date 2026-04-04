# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

# Title
t1 = doc.add_heading('交通银行安徽省分行', level=0)
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2 = doc.add_heading('AI兴趣小组第一期培训班实施方案', level=0)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('编制单位：交通银行安徽省分行\n').bold = True
meta.add_run('编制日期：' + datetime.datetime.now().strftime('%Y年%m月%d日'))
doc.add_paragraph()

# === Section 1 ===
doc.add_heading('一、背景与目的', level=1)
doc.add_heading('1.1 背景', level=2)
doc.add_paragraph('交通银行安徽省分行于2026年3月发起成立"AI兴趣小组"，旨在搭建全行员工AI应用学习与实践平台。小组成立至今，群成员已达200人，具备良好的群众基础。为推动AI工具在日常工作中的落地应用，现组织第一期线下培训班。')

doc.add_heading('1.2 调研支撑', level=2)
doc.add_paragraph('基于59份调研问卷分析：')
for bold_text, normal_text in [
    ('参与意愿极强：', '81%成员表示希望加入学习，80%希望系统学习AI'),
    ('核心痛点需求：', '写材料（19人）、数据分析（18人）、减少重复工作（16人）'),
    ('最大使用障碍：', '信息安全顾虑（18人）、不会提问/提示词（15人）、AI不够专业（15人）'),
    ('首选学习形式：', 'AI工具实操课，工作日晚上线上参与'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bold_text).bold = True
    p.add_run(normal_text)

doc.add_heading('1.3 培训目标', level=2)
for g in [
    '让学员掌握交心AI及AI办公助手的基本操作，能够应用于日常文字材料撰写与数据分析',
    '建立AI应用的安全意识，了解合规使用规范，消除信息安全顾虑',
    '通过分组研讨，梳理各条线业务AI落地思路，形成可执行的应用方案',
    '培训结束后持续跟踪，通过作业机制推动AI工具真正落地使用',
]:
    doc.add_paragraph(g, style='List Bullet')
doc.add_paragraph()

# === Section 2 ===
doc.add_heading('二、组织架构', level=1)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['角色', '主要职责', '建议人选']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for i, row_data in enumerate([
    ['班主任', '统筹培训整体安排，协调资源，处理突发问题', '分行人力资源部/金融科技部负责人'],
    ['执行负责人', '课程执行、物料准备、学员签到、现场调度', 'AI兴趣小组发起人（Leon）'],
    ['讲师团', '课程设计与授课（交心AI、AI办公助手、实操练习）', '金融科技部骨干 + 外聘专家'],
    ['技术支持', '设备调试、网络保障、线上直播平台运维', '金融科技部IT支持'],
]):
    for j, text in enumerate(row_data):
        table.rows[i+1].cells[j].text = text
doc.add_paragraph()

# === Section 3 ===
doc.add_heading('三、培训对象与名额分配', level=1)
doc.add_heading('3.1 线下培训（主场地）', level=2)
doc.add_paragraph('建议各市辖行选派2-3名业务骨干参加线下培训，名额分配如下：')
cities = ['合肥', '芜湖', '蚌埠', '淮南', '马鞍山', '淮北', '铜陵', '安庆', '黄山', '阜阳', '宿州', '滁州', '六安', '宣城', '池州', '亳州']
t = doc.add_table(rows=len(cities)+1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = '辖行'
t.rows[0].cells[1].text = '建议名额'
t.rows[0].cells[0].paragraphs[0].runs[0].bold = True
t.rows[0].cells[1].paragraphs[0].runs[0].bold = True
for i, city in enumerate(cities):
    t.rows[i+1].cells[0].text = city + '分行'
    t.rows[i+1].cells[1].text = '2-3人'
p = doc.add_paragraph()
p.add_run('线下主场地总容量：约40-50人').bold = True

doc.add_heading('3.2 线上参会', level=2)
doc.add_paragraph('其余报名参培训的学员可通过线上方式参会，全程观看直播并参与线上互动答疑。线上学员可参与课后作业提交，但优秀作业评选线下学员优先。')

doc.add_heading('3.3 报名方式', level=2)
doc.add_paragraph('由各市辖行HR或科技专管员汇总本地报名名单，统一通过邮件或内部问卷系统提交，报名截止日期建议为培训举办前10个工作日。')
doc.add_paragraph()

# === Section 4 ===
doc.add_heading('四、培训形式', level=1)
t4 = doc.add_table(rows=6, cols=2)
t4.style = 'Table Grid'
t4.rows[0].cells[0].text = '项目'
t4.rows[0].cells[1].text = '说明'
t4.rows[0].cells[0].paragraphs[0].runs[0].bold = True
t4.rows[0].cells[1].paragraphs[0].runs[0].bold = True
for i, (k, v) in enumerate([
    ('培训时间', '建议选取周四-周五（2天），便于外地学员往返'),
    ('培训地点', '交通银行安徽省分行（具体会议室待定）'),
    ('培训形式', '线下集中授课 + 线上同步直播，线下学员实操为主，线上学员观摩互动'),
    ('设备需求', '主屏投影、分组讨论桌椅、有线/无线网络保障、直播设备（腾讯会议/飞书会议）'),
    ('封闭网络说明', '本次培训基于封闭网络环境，课程内容锚定行内已部署工具（交心AI、AI办公助手），确保学员学完即可使用'),
]):
    t4.rows[i+1].cells[0].text = k
    t4.rows[i+1].cells[1].text = v
doc.add_paragraph()

# === Section 5 ===
doc.add_heading('五、课程安排', level=1)
p5 = doc.add_paragraph('核心工具：交心AI（类豆包）、AI办公助手（行内部署）')
p5.italic = True

doc.add_heading('Day 1（第一天）——AI应用基础与核心场景', level=2)
day1 = [
    ('08:30-09:00', '报到签到', '领取物料，熟悉环境', '班主任团队'),
    ('09:00-09:30', '开班仪式', '领导致辞 + 培训目标说明 + 破冰（学员自我介绍，每人1分钟）', '班主任'),
    ('09:30-10:30', '交心AI：提示词工程', '结构化提示词方法（角色、任务、格式、约束）；常用技巧演示；学员现场练习', '讲师'),
    ('10:30-10:45', '茶歇', '', ''),
    ('10:45-12:00', '交心AI：场景练习', '结合银行业务场景（写材料、改稿、润色、提炼摘要）；学员动手实操', '讲师'),
    ('12:00-14:00', '午餐 + 休息', '', ''),
    ('14:00-15:30', 'AI办公助手：基础操作', '功能介绍 + 常用场景演示（会议纪要生成、报告辅助撰写、表格整理）', '讲师'),
    ('15:30-15:45', '茶歇', '', ''),
    ('15:45-17:30', 'AI办公助手：实操练习', '学员分组使用AI办公助手完成预设任务，讲师巡回指导', '讲师 + 技术支持'),
]
t5 = doc.add_table(rows=len(day1)+1, cols=4)
t5.style = 'Table Grid'
for j, h in enumerate(['时间', '内容', '详情', '负责人']):
    t5.rows[0].cells[j].text = h
    t5.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, row in enumerate(day1):
    for j, val in enumerate(row):
        t5.rows[i+1].cells[j].text = val
doc.add_paragraph()

doc.add_heading('Day 2（第二天）——进阶应用、研讨与展望', level=2)
day2 = [
    ('09:00-09:30', '前一天内容回顾', '答疑 + 学员分享使用心得（3-5人）', '班主任'),
    ('09:30-10:45', 'AI辅助数据分析', '多光谱数据/业务数据AI分析入门；工具演示 + 动手练习', '讲师'),
    ('10:45-11:00', '茶歇', '', ''),
    ('11:00-12:00', 'AI绘图与辅助汇报', 'PPT辅助生成、图片素材AI制作（基于行内工具或外部演示）', '讲师'),
    ('12:00-14:00', '午餐 + 休息', '', ''),
    ('14:00-15:15', '分组研讨', '结合各条线业务，梳理AI落地思路，每组产出行级应用建议（3-5条）', '各组召集人'),
    ('15:15-15:30', '茶歇', '', ''),
    ('15:30-16:15', 'OpenClaw与QClaw介绍', '介绍多Agent框架与AI发展趋势，建立技术自信；QClaw演示', '讲师'),
    ('16:15-16:45', '作业布置与结班', '公布作业要求、提交方式、截止时间、评优机制；领导总结', '班主任'),
]
t6 = doc.add_table(rows=len(day2)+1, cols=4)
t6.style = 'Table Grid'
for j, h in enumerate(['时间', '内容', '详情', '负责人']):
    t6.rows[0].cells[j].text = h
    t6.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, row in enumerate(day2):
    for j, val in enumerate(row):
        t6.rows[i+1].cells[j].text = val
doc.add_paragraph()

# === Section 6 ===
doc.add_heading('六、培训材料清单', level=1)
materials = [
    ('学员手册', '含课程大纲、每日笔记页、AI工具快速上手指南', '纸质版或PDF'),
    ('交心AI使用指南', '基础操作手册 + 常用提示词模板库', '纸质版或PDF'),
    ('AI办公助手操作手册', '各功能模块说明 + 银行业务场景示例', '纸质版或PDF'),
    ('AI使用安全指南', '信息安全合规须知，消除学员使用顾虑', '纸质版，人手一份'),
    ('课前预习材料', '可选：AI基础概念短视频（3-5分钟），发至学员群', '线上发送'),
    ('作业模板', '作业提交表格式模板，含各条线业务场景提示', '线上发送'),
]
t7 = doc.add_table(rows=len(materials)+1, cols=3)
t7.style = 'Table Grid'
for j, h in enumerate(['材料名称', '内容说明', '格式']):
    t7.rows[0].cells[j].text = h
    t7.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, row in enumerate(materials):
    for j, val in enumerate(row):
        t7.rows[i+1].cells[j].text = val
doc.add_paragraph()

# === Section 7 ===
doc.add_heading('七、宣传与通知', level=1)
notif = [
    ('通知时间', '培训举办前15个工作日发出第一轮通知，培训前7个工作日发出第二轮提醒'),
    ('通知渠道', '省分行邮件系统 + 内部工作群 + AI兴趣小组群公告'),
    ('报名方式', '各市辖行统一汇总后提交报名回执（邮件或内部系统）'),
    ('温馨提示', '告知封闭网络环境、自带笔记本（如需）、课前安装相关工具等准备事项'),
]
t8 = doc.add_table(rows=len(notif)+1, cols=2)
t8.style = 'Table Grid'
t8.rows[0].cells[0].text = '项目'
t8.rows[0].cells[1].text = '说明'
t8.rows[0].cells[0].paragraphs[0].runs[0].bold = True
t8.rows[0].cells[1].paragraphs[0].runs[0].bold = True
for i, (k, v) in enumerate(notif):
    t8.rows[i+1].cells[0].text = k
    t8.rows[i+1].cells[1].text = v
doc.add_paragraph()

# === Section 8 ===
doc.add_heading('八、培训后跟踪机制', level=1)
doc.add_heading('8.1 作业设计', level=2)
doc.add_paragraph('作业结合学员本职岗位布置，建议格式：')
for h in [
    '任务描述：用AI工具（交心AI或AI办公助手）辅助完成一件本职工作相关任务',
    '输出物：提交任务完成结果（文档、报告、数据分析等） + 简要说明使用了哪些AI功能',
    '字数要求：300-500字',
]:
    doc.add_paragraph(h, style='List Bullet')

doc.add_heading('8.2 提交与评优', level=2)
submit = [
    ('提交截止', '培训结束后10个工作日'),
    ('提交方式', '邮件或内部系统，文件名格式：姓名-机构-岗位-AI作业'),
    ('评优机制', '优秀作业在兴趣小组群内展示，讲师点评，颁发纪念证书/奖品'),
    ('后续跟踪', '优秀作业可纳入分行AI应用案例库，供全行学习借鉴'),
]
t9 = doc.add_table(rows=len(submit)+1, cols=2)
t9.style = 'Table Grid'
t9.rows[0].cells[0].text = '项目'
t9.rows[0].cells[1].text = '说明'
t9.rows[0].cells[0].paragraphs[0].runs[0].bold = True
t9.rows[0].cells[1].paragraphs[0].runs[0].bold = True
for i, (k, v) in enumerate(submit):
    t9.rows[i+1].cells[0].text = k
    t9.rows[i+1].cells[1].text = v

doc.add_heading('8.3 持续答疑与打卡', level=2)
doc.add_paragraph('培训结束后，兴趣小组群持续运营，讲师或执行负责人在群内定期答疑，鼓励学员分享使用心得。后续可每月组织一次线上答疑会或AI使用分享会，形成持续学习氛围。')
doc.add_paragraph()

# === Section 9 ===
doc.add_heading('九、预算估算（参考）', level=1)
budget = [
    ('场地及设备', '布置费用（背景板、条幅等）', '待定'),
    ('物料印刷', '学员手册、AI使用指南、安全指南等（按50人份）', '待定'),
    ('茶歇', '两天上下午茶歇（咖啡、茶点）', '待定'),
    ('讲师费用', '外聘专家费用（如有）', '待定'),
    ('奖品/证书', '优秀作业奖品、结业证书', '待定'),
    ('其他', '直播设备耗材、线上平台费用等', '待定'),
]
t10 = doc.add_table(rows=len(budget)+1, cols=3)
t10.style = 'Table Grid'
for j, h in enumerate(['类别', '明细', '预算']):
    t10.rows[0].cells[j].text = h
    t10.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, row in enumerate(budget):
    for j, val in enumerate(row):
        t10.rows[i+1].cells[j].text = val
doc.add_paragraph()

# === Section 10 ===
doc.add_heading('十、风险与预案', level=1)
risks = [
    ('网络/系统故障', '培训场地封闭网络或AI工具无法访问', '提前测试，确认工具可用；准备降级方案（如本地演示）'),
    ('人数超预期', '线下报名人数超过场地容量', '优先保障地市行骨干；线上通道开放不受限制'),
    ('参训学员积极性不足', '报名后参与度低', '破冰环节设计；作业与评优机制激励；班主任跟进'),
    ('讲师时间冲突', '预设讲师无法按期授课', '准备备用讲师或调整课程顺序'),
]
t11 = doc.add_table(rows=len(risks)+1, cols=3)
t11.style = 'Table Grid'
for j, h in enumerate(['风险场景', '描述', '应对预案']):
    t11.rows[0].cells[j].text = h
    t11.rows[0].cells[j].paragraphs[0].runs[0].bold = True
for i, row in enumerate(risks):
    for j, val in enumerate(row):
        t11.rows[i+1].cells[j].text = val
doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('—— 交通银行安徽省分行 AI兴趣小组 ——')

output = r'C:\Users\flyan\.openclaw\workspace\agents-workspaces\secretary\projects\AI兴趣小组\AI兴趣小组第一期培训班实施方案.docx'
doc.save(output)
print('Saved:', output)
